#!/usr/bin/env python3
"""Génère le dossier CESI Bloc 3 en réutilisant le modèle Word CESIZen v3."""

from __future__ import annotations

import io
import os
import sys
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/home/harwain/Téléchargements/Documentation_Technique_CESIZen_v3.docx")
OUTPUT = ROOT / "livrables" / "Dossier_Bloc_3_Deployement_Maintenance_Securite_CESIZen.docx"

BLUE = "2E75B6"
DARK_BLUE = "203864"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F2F2F2"
WHITE = "FFFFFF"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "F4CCCC"


def clear_body(doc: Document) -> None:
    body = doc._element.body
    section = body.sectPr
    for child in list(body):
        if child is not section:
            body.remove(child)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def style_by_id(doc, style_id: str):
    return next(style for style in doc.styles if style.style_id == style_id)


def add_para(doc, text="", *, bold=False, italic=False, align=None, size=None, color=None,
             before=0, after=4, line=1.0, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc, items, *, level=0, size=9.2, compact=True):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.55 + 0.45 * level)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(1.5 if compact else 3)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"• {item}")
        r.font.size = Pt(size)


def add_numbered(doc, items, *, size=9.2):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(f"{index}. {item}").font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=style_by_id(doc, str(932 + level)))
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths=None, font_size=7.6, header_fill=BLUE,
              first_col_fill=None, align=WD_TABLE_ALIGNMENT.CENTER):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style_by_id(doc, "751")
    table.alignment = align
    table.autofit = False
    hdr = table.rows[0]
    repeat_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, header_fill)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(font_size)
        if widths:
            cell.width = Cm(widths[i])
    for ri, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if first_col_fill and i == 0:
                shade(cells[i], first_col_fill)
            elif ri % 2 == 1:
                shade(cells[i], "F8FBFD")
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 0.95
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
            if i == 0 and first_col_fill:
                r.bold = True
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    doc.add_page_break()


def update_headers_footers(doc):
    for section in doc.sections:
        for p in section.header.paragraphs:
            for run in p.runs:
                if "Documentation Technique" in run.text:
                    run.text = "Bloc 3 — Déploiement, maintenance et sécurité — CESIZen"
        for p in section.footer.paragraphs:
            for run in p.runs:
                if "Marin Cadro" in run.text:
                    run.text = run.text.replace("17/05/2026", "13/08/2026")


def add_cover(doc, logo_cesi, logo_ministry):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    logos = doc.add_table(rows=1, cols=2)
    logos.alignment = WD_TABLE_ALIGNMENT.CENTER
    logos.autofit = False
    logos.columns[0].width = Cm(8)
    logos.columns[1].width = Cm(8)
    p = logos.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(logo_ministry, width=Cm(4.0))
    p = logos.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(logo_cesi, width=Cm(4.3))

    add_para(doc, "", after=32)
    add_para(doc, "DÉPLOYER ET SÉCURISER", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             size=27, color=DARK_BLUE, after=0)
    add_para(doc, "LES APPLICATIONS INFORMATIQUES", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=27, color=DARK_BLUE, after=14)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)
    add_para(doc, "CESIZen", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             size=22, color=BLUE, after=4)
    add_para(doc, "Plan de déploiement, de maintenance et de sécurisation",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, after=50)

    metadata = [
        ("Auteur", "Marin Cadro"),
        ("Date", "13/08/2026"),
        ("Version", "1.1"),
        ("Formation", "Concepteur développeur d’applications — Bloc 3"),
        ("Dépôt", "https://github.com/MarinCD/CesiZen"),
        ("Recette", "https://cesizen.optihent.fr"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = style_by_id(doc, "751")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(metadata):
        table.cell(i, 0).width = Cm(4.0)
        table.cell(i, 1).width = Cm(10.0)
        shade(table.cell(i, 0), BLUE)
        for cell in table.rows[i].cells:
            margins(cell, 60, 90, 60, 90)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        lr = table.cell(i, 0).paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.color.rgb = RGBColor(255, 255, 255)
        lr.font.size = Pt(9)
        vr = table.cell(i, 1).paragraphs[0].add_run(value)
        vr.font.size = Pt(9)
    page_break(doc)


def add_toc(doc):
    add_heading(doc, "Sommaire", 1)
    entries = [
        ("1. Introduction et périmètre", 3),
        ("2. Architecture et contexte", 4),
        ("3. Environnements de déploiement", 5),
        ("4. Plan de déploiement", 6),
        ("5. Versioning et intégration continue", 7),
        ("6. Ressources, exploitation et retour arrière", 8),
        ("7. Outil de maintenance et ticketing", 9),
        ("8. Méthodologie de maintenance", 10),
        ("9. Veille technologique", 11),
        ("10. Stratégie de sécurisation", 12),
        ("11. Matrice des risques et plan d’actions", 13),
        ("12. Données personnelles et RGPD", 15),
        ("13. Gestion de crise", 16),
        ("14. Bonnes pratiques et validation", 17),
        ("15. Démonstration et éléments de preuve", 18),
        ("16. Limites, évolutions et conclusion", 19),
        ("17. Glossaire et références", 20),
    ]
    for title, number in entries:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(14.6)
        table.columns[1].width = Cm(1.0)
        p = table.cell(0, 0).paragraphs[0]
        p.paragraph_format.space_after = Pt(1.5)
        r = p.add_run(title)
        r.font.size = Pt(10.5)
        p = table.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(str(number)).font.size = Pt(10.5)
    add_para(doc, "Lecture du statut", bold=True, color=DARK_BLUE, before=14, after=5)
    add_table(doc, ["Statut", "Signification"], [
        ("En place", "Mécanisme implémenté et vérifié dans le dépôt ou l’environnement de recette."),
        ("Procédure", "Organisation définie pour une future exploitation contrôlée."),
        ("Limite", "Écart connu, explicitement accepté ou planifié."),
    ], widths=[3.2, 12.4], font_size=8.5)
    page_break(doc)


def build_document():
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Modèle introuvable : {TEMPLATE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    update_headers_footers(doc)

    with zipfile.ZipFile(TEMPLATE) as archive:
        logo_cesi = io.BytesIO(archive.read("word/media/image1.png"))
        logo_ministry = io.BytesIO(archive.read("word/media/image2.png"))

    add_cover(doc, logo_cesi, logo_ministry)
    add_toc(doc)

    # Page 3
    add_heading(doc, "1. Introduction et périmètre", 1)
    add_heading(doc, "1.1. Contexte", 2)
    add_para(doc, "CESIZen est une application web de prévention et d’accompagnement autour de la santé mentale. Elle propose des contenus d’information, un questionnaire de stress inspiré de l’échelle Holmes & Rahe, un historique pour les comptes autorisés et un back-office. Le présent dossier répond au Bloc 3 du titre Concepteur développeur d’applications : déployer, maintenir et sécuriser une application informatique.")
    add_heading(doc, "1.2. Objectifs du dossier", 2)
    add_bullets(doc, [
        "Décrire une architecture de déploiement externalisée, dimensionnée et reproductible.",
        "Présenter les environnements de développement, d’intégration et de recette, dont un est réellement configuré.",
        "Démontrer l’usage de GitHub, GitHub Actions, Dokploy, Cloudflare et Prisma Migrate.",
        "Organiser le traitement des anomalies et évolutions avec GitHub Issues et une méthode traçable.",
        "Analyser les risques, leur criticité et les mesures correctives ou préventives.",
        "Définir la gestion des données personnelles, la continuité et la réponse à incident.",
    ], size=9.4)
    add_heading(doc, "1.3. Périmètre et positionnement", 2)
    add_para(doc, "Le périmètre fonctionnel retenu couvre les modules obligatoires Comptes et Informations, ainsi que le module au choix Diagnostics. L’environnement publié à l’adresse cesizen.optihent.fr est un environnement de recette et de démonstration. Il utilise un mode d’exécution optimisé de production, mais n’est pas qualifié pour conserver des données de santé : l’historisation des diagnostics y reste désactivée tant qu’une preuve HDS n’est pas disponible.")
    add_table(doc, ["Élément", "État vérifié au 13/08/2026"], [
        ("Application", "Next.js 16.3.0, React 19.2.8, NextAuth 4.24.15"),
        ("Qualité", "180 tests réussis ; typecheck, lint et build validés"),
        ("Sécurité", "npm audit : 0 vulnérabilité connue sur 716 dépendances"),
        ("Recette", "HTTPS actif ; HTTP redirigé en 308 ; réponse publique 200"),
    ], widths=[3.2, 12.4], font_size=8.4)
    page_break(doc)

    # Page 4
    add_heading(doc, "2. Architecture et contexte", 1)
    add_heading(doc, "2.1. Architecture logique", 2)
    add_table(doc, ["Flux", "Composant", "Responsabilité"], [
        ("1", "Navigateur", "Interface accessible, session utilisateur et appels HTTPS."),
        ("2", "Cloudflare DNS", "Résolution de cesizen.optihent.fr et protection du point d’entrée."),
        ("3", "Traefik / Dokploy", "Terminaison TLS, routage par nom d’hôte et orchestration du conteneur."),
        ("4", "Next.js", "Rendu serveur, routes API, authentification, validation et règles métier."),
        ("5", "Prisma ORM", "Accès paramétré, modèles relationnels et migrations versionnées."),
        ("6", "MySQL Alwaysdata", "Persistance des comptes, contenus, diagnostics autorisés et audits."),
        ("7", "GitHub", "Sources, tickets, historique, Dependabot et intégration continue."),
    ], widths=[1.0, 4.0, 10.6], font_size=8.1)
    add_heading(doc, "2.2. Flux de publication", 2)
    add_table(doc, ["Développeur", "GitHub", "CI", "Dokploy", "Utilisateur"], [[
        "Branche de travail", "develop → main", "audit + lint + types + tests + build", "main déployée", "HTTPS"
    ]], widths=[3.0, 3.0, 4.3, 3.1, 2.2], font_size=7.6)
    add_heading(doc, "2.3. Choix structurants", 2)
    add_bullets(doc, [
        "Monolithe Next.js : périmètre individuel, déploiement simple et types partagés client/serveur.",
        "MySQL managé séparément : données indépendantes du cycle de vie du conteneur.",
        "Dokploy et Traefik : plusieurs projets partagent les ports 80/443 sans exposer le port 3000.",
        "Cloudflare : gestion centralisée de la zone DNS ; les autres sous-domaines restent isolés.",
        "Prisma : schéma source de vérité, requêtes paramétrées et migration reproductible.",
    ], size=9.2)
    add_heading(doc, "2.4. Contraintes", 2)
    add_para(doc, "Le projet est réalisé individuellement, avec un trafic de démonstration et un budget limité. La disponibilité dépend du VPS, de Cloudflare et d’Alwaysdata. Les résultats de diagnostic peuvent être assimilés à des données de santé : l’environnement actuel calcule le score sans le conserver. Cette limite est assumée et contrôlée par HDS_COMPLIANT_STORAGE=0.")
    page_break(doc)

    # Page 5
    add_heading(doc, "3. Environnements de déploiement", 1)
    add_table(doc, ["Environnement", "Usage", "Configuration", "Données", "Déploiement"], [
        ("Développement", "Coder et déboguer", "Poste local, Node 22, npm, Next dev", "Base dédiée ou mocks", "Manuel, branche de travail"),
        ("Intégration CI", "Valider chaque changement", "GitHub Actions, Ubuntu, MySQL 8.4", "Base éphémère", "Automatique sur push/PR"),
        ("Recette", "Démonstration client/CESI", "VPS, Dokploy, Traefik, Cloudflare", "Alwaysdata ; santé non historisée", "Promotion develop → main"),
        ("Production cible", "Usage réel", "Hébergeur qualifié, supervision et astreinte", "Stockage validé HDS si nécessaire", "Release signée et fenêtre contrôlée"),
    ], widths=[2.4, 3.0, 4.5, 3.4, 2.5], font_size=7.5)
    add_heading(doc, "3.1. Environnement réellement configuré", 2)
    add_para(doc, "L’environnement de recette est réellement accessible sur https://cesizen.optihent.fr. Le DNS pointe vers le VPS, Traefik sélectionne le service grâce au nom d’hôte, et Let’s Encrypt fournit le certificat utilisé entre le navigateur et l’origine. Cloudflare est configuré en mode Full (strict). Le conteneur écoute uniquement sur le port interne 3000.")
    add_heading(doc, "3.2. Variables et séparation", 2)
    add_table(doc, ["Variable", "Rôle", "Règle"], [
        ("DATABASE_URL", "Connexion MySQL", "Secret Dokploy, jamais commité"),
        ("NEXTAUTH_URL", "Origine canonique", "URL HTTPS de l’environnement"),
        ("NEXTAUTH_SECRET", "Signature de session", "Aléatoire et distinct"),
        ("AUDIT_HMAC_KEY", "Pseudonymisation audit", "Distinct du secret de session"),
        ("RATE_LIMIT_HMAC_KEY", "Clé des compteurs", "Distinct et renouvelable"),
        ("HDS_COMPLIANT_STORAGE", "Barrière santé", "0 sans validation HDS"),
    ], widths=[3.7, 5.0, 6.9], font_size=7.8)
    add_para(doc, "Les secrets sont injectés au niveau du service Dokploy et masqués dans les preuves. Le fichier .env reste exclu de Git. Une rotation implique la mise à jour du coffre, un redéploiement puis la révocation de l’ancienne valeur.", italic=True, size=8.8)
    page_break(doc)

    # Page 6
    add_heading(doc, "4. Plan de déploiement", 1)
    add_heading(doc, "4.1. Préparation", 2)
    add_numbered(doc, [
        "Créer une branche feature/fix depuis develop, lier le ticket et définir les critères d’acceptation.",
        "Contrôler npm audit, lint, typecheck, tests et build ; interdire toute vulnérabilité élevée ou critique non acceptée.",
        "Vérifier la sauvegarde de la base et la procédure de retour arrière.",
        "Relire la migration Prisma ; séparer les évolutions compatibles des suppressions destructives.",
        "Fusionner la pull request vers develop après CI, réaliser la recette, puis promouvoir develop vers main par une seconde pull request.",
        "Valider les variables Dokploy sans exposer leur valeur dans les captures ou tickets.",
    ], size=8.8)
    add_heading(doc, "4.2. Exécution en recette", 2)
    add_table(doc, ["Étape", "Action", "Contrôle / preuve"], [
        ("1", "GitHub récupère la révision de main", "SHA du commit et auteur"),
        ("2", "Nixpacks exécute npm ci", "Lockfile respecté ; installation reproductible"),
        ("3", "Prisma Client est généré", "Schéma parsé sans erreur"),
        ("4", "Next.js produit le build", "Compilation et types réussis"),
        ("5", "Prisma migrate deploy", "Baseline connue ; schéma à jour"),
        ("6", "Next.js écoute sur 0.0.0.0:3000", "Service stable, aucune boucle de restart"),
        ("7", "Traefik publie le domaine", "HTTPS 200 et HTTP 308"),
        ("8", "Smoke tests", "Accueil, connexion, informations, diagnostic anonyme, admin"),
    ], widths=[1.2, 7.2, 7.2], font_size=8.0)
    add_heading(doc, "4.3. Décision de mise à disposition", 2)
    add_para(doc, "La version est acceptée si la CI est verte, que la migration est réussie, qu’aucune erreur critique n’apparaît dans les logs et que les smoke tests sont validés. En cas d’échec, la publication est interrompue et le dernier artefact sain est conservé. La communication indique la version, la fenêtre, l’impact, le responsable et l’issue de validation.")
    page_break(doc)

    # Page 7
    add_heading(doc, "5. Versioning et intégration continue", 1)
    add_heading(doc, "5.1. Gestion des sources", 2)
    add_para(doc, "GitHub centralise le dépôt MarinCD/CesiZen. Deux branches permanentes sont configurées : develop agrège les changements validés pour l’intégration et main contient uniquement la version stable déployable par Dokploy. Les branches de travail sont créées depuis develop et y reviennent par pull request. Après recette, une pull request develop vers main matérialise la promotion. Les migrations, le package-lock et la documentation sont versionnés ; les secrets, sauvegardes, node_modules et sorties de build sont ignorés.")
    add_table(doc, ["Objet", "Convention", "Finalité"], [
        ("develop", "Branche d’intégration", "Cible des PR feature/fix/docs"),
        ("main", "Branche stable", "Source du déploiement Dokploy"),
        ("Branche courte", "feat/…, fix/…, docs/…", "Isoler un changement depuis develop"),
        ("hotfix", "hotfix/… depuis main", "Correction urgente reportée aussi sur develop"),
        ("Commit", "type: description concise", "Historique lisible"),
        ("Pull request", "issue, risques, tests, rollback", "Revue et décision"),
        ("Version", "SemVer vMAJEUR.MINEUR.CORRECTIF", "Identifier un déploiement"),
        ("Migration", "horodatage_description", "Ordre déterministe"),
    ], widths=[3.0, 5.3, 7.3], font_size=8.2)
    add_heading(doc, "5.2. Pipeline GitHub Actions", 2)
    add_table(doc, ["Contrôle", "But", "Blocage"], [
        ("npm ci", "Dépendances conformes au lockfile", "Échec d’installation"),
        ("prisma migrate deploy", "Tester la migration sur MySQL éphémère", "Migration impossible"),
        ("npm audit --audit-level=high", "Détecter les failles connues", "Vulnérabilité haute/critique"),
        ("ESLint + TypeScript", "Qualité statique", "Erreur de lint ou de types"),
        ("Vitest", "180 tests unitaires/intégration", "Régression"),
        ("next build", "Valider l’artefact", "Compilation impossible"),
    ], widths=[4.0, 7.0, 4.6], font_size=8.0)
    add_heading(doc, "5.3. Pilotage", 2)
    add_para(doc, "Le run GitHub Actions du commit 9ce6e4e est terminé avec succès. La CI se déclenche sur les pull requests et les pushes de main et develop. Les deux branches sont protégées : pull request obligatoire, contrôle quality-security-build à jour, résolution des conversations, historique linéaire, suppression et force-push interdits. Aucune approbation d’un second compte n’est imposée, car l’évaluation est individuelle. Dependabot surveille npm et GitHub Actions.")
    page_break(doc)

    # Page 8
    add_heading(doc, "6. Ressources, exploitation et retour arrière", 1)
    add_heading(doc, "6.1. Dimensionnement", 2)
    add_table(doc, ["Ressource", "Recette proposée", "Surveillance / seuil"], [
        ("CPU", "2 vCPU", "Alerte si > 80 % pendant 10 min"),
        ("Mémoire", "4 Go + swap contrôlée", "Alerte > 85 % ; exit 137 au build"),
        ("Disque", "40 Go minimum", "Alerte > 80 % ; purge images/logs"),
        ("Réseau", "80/443 publics ; 3000 interne", "Disponibilité HTTPS et latence"),
        ("MySQL", "Service Alwaysdata séparé", "Connexions, espace, sauvegarde"),
        ("Sauvegarde", "Quotidienne ; conservation 30 jours", "RPO 24 h, RTO 4 h"),
    ], widths=[3.0, 6.4, 6.2], font_size=8.0)
    add_heading(doc, "6.2. Supervision minimale", 2)
    add_bullets(doc, [
        "Disponibilité de la page d’accueil et validité du certificat TLS.",
        "État et redémarrages du conteneur dans Dokploy ; consommation CPU/RAM/disque.",
        "Erreurs serveur, échecs d’authentification, rate-limit et événements d’audit pseudonymisés.",
        "Échec des CI, alertes Dependabot et dérive du schéma Prisma.",
    ], size=9.0)
    add_heading(doc, "6.3. Retour arrière", 2)
    add_numbered(doc, [
        "Geler les écritures si une donnée ou migration est concernée et ouvrir un incident.",
        "Redéployer l’image ou le commit sain identifié ; ne pas réécrire l’historique Git.",
        "Pour une migration compatible, laisser les colonnes ajoutées jusqu’au correctif suivant.",
        "Pour une restauration, créer une base isolée, déchiffrer la sauvegarde, restaurer et contrôler les volumes.",
        "Exécuter migrate deploy, smoke tests et contrôles fonctionnels avant bascule.",
        "Documenter les temps réels, l’impact, la décision et le retour d’expérience.",
    ], size=8.7)
    add_para(doc, "État : deux sauvegardes applicatives chiffrées ont été authentifiées avant et après la baseline. Un test complet de restauration sur une base temporaire reste une opération mensuelle à réaliser.", italic=True, size=8.7)
    page_break(doc)

    # Page 9
    add_heading(doc, "7. Outil de maintenance et ticketing", 1)
    add_heading(doc, "7.1. Choix de GitHub Issues", 2)
    add_para(doc, "GitHub Issues est cohérent avec le dépôt et le dimensionnement individuel. Il relie chaque demande aux commits, pull requests, versions et résultats de CI. Des modèles dédiés encadrent les anomalies et maintenances ; les formulaires rappellent de ne joindre ni secret, ni adresse IP, ni donnée de santé.")
    add_table(doc, ["Configuration", "Mise en œuvre"], [
        ("Types", "Anomalie et Maintenance"),
        ("Labels", "bug, maintenance, security, dependencies, priority:P1, triage"),
        ("Milestones", "Architecture, Comptes, Informations, Diagnostics, Sécurité/RGPD, Tests/Déploiement"),
        ("États projet", "Backlog → Qualifié → En cours → En revue → Validé → Déployé"),
        ("PR", "Objet, issue, risques, migration, retour arrière et preuves"),
        ("Confidentiel", "GitHub Private Vulnerability Reporting"),
    ], widths=[4.0, 11.6], font_size=8.3)
    add_heading(doc, "7.2. Exemple réel", 2)
    add_para(doc, "L’issue #32 « Corriger les dépendances npm vulnérables » qualifie une priorité P1, liste les avis GHSA, définit les critères d’acceptation et le retour arrière. La correction a conduit à Next.js 16.3.0, NextAuth 4.24.15 et un audit sans vulnérabilité connue. L’issue #31 suit la publication de la plateforme.")
    add_heading(doc, "7.3. Informations obligatoires d’un ticket", 2)
    add_bullets(doc, [
        "Contexte, environnement, résultat observé, résultat attendu et reproduction.",
        "Impact métier, sécurité, données concernées et priorité P1 à P4.",
        "Responsable, échéance, critères d’acceptation et dépendances.",
        "Plan de test, sauvegarde, procédure de retour arrière et preuves expurgées.",
        "Version corrigée, date de déploiement, validation et motif de clôture.",
    ], size=9.0)
    page_break(doc)

    # Page 10
    add_heading(doc, "8. Méthodologie de maintenance", 1)
    add_heading(doc, "8.1. Cycle de traitement", 2)
    add_table(doc, ["Phase", "Activité", "Sortie"], [
        ("Détection", "Utilisateur, monitoring, test, Dependabot ou veille", "Ticket horodaté"),
        ("Qualification", "Reproduction, périmètre, données, criticité", "Priorité et responsable"),
        ("Planification", "Solution, charge, dépendances, retour arrière", "Critères d’acceptation"),
        ("Correction", "Branche dédiée et changement minimal", "Commit lié"),
        ("Validation", "Revue, CI, tests de confirmation/régression", "Preuves"),
        ("Déploiement", "Sauvegarde, migration, smoke tests", "Version disponible"),
        ("Clôture", "Validation, documentation et mesure", "Ticket clos / REX"),
    ], widths=[2.7, 8.0, 4.9], font_size=8.1)
    add_heading(doc, "8.2. Priorités et engagements", 2)
    add_table(doc, ["Priorité", "Critère", "Prise en charge", "Cible"], [
        ("P1", "Compromission, fuite, indisponibilité", "1 heure", "24 heures"),
        ("P2", "Fonction majeure ou faille exploitable", "1 jour ouvré", "7 jours"),
        ("P3", "Impact limité, contournement possible", "3 jours ouvrés", "30 jours"),
        ("P4", "Dette, documentation, confort", "Prochain tri", "Cycle planifié"),
    ], widths=[2.0, 7.1, 3.5, 3.0], font_size=8.0)
    add_heading(doc, "8.3. Rôles", 2)
    add_para(doc, "Le responsable applicatif priorise et accepte ; le développeur corrige et teste ; le relecteur contrôle le code et les preuves ; l’exploitant sauvegarde, publie et surveille. Dans le projet individuel, Marin Cadro tient ces rôles successivement : leur séparation est matérialisée par les statuts, les contrôles CI et les preuves du ticket.")
    add_heading(doc, "8.4. Mesures", 2)
    add_bullets(doc, [
        "Délai moyen de prise en charge et de correction ; respect du SLA.",
        "Taux de réouverture, régressions et changements retournés en arrière.",
        "Vulnérabilités ouvertes par criticité et âge des dépendances.",
        "Succès des sauvegardes, restaurations et déploiements.",
    ], size=8.8)
    page_break(doc)

    # Page 11
    add_heading(doc, "9. Veille technologique", 1)
    add_heading(doc, "9.1. Méthode", 2)
    add_para(doc, "La veille suit un cycle court : collecter des sources officielles, filtrer selon la stack, qualifier l’impact, expérimenter hors production, ouvrir un ticket puis capitaliser la décision. Un créneau hebdomadaire traite les alertes ; une synthèse mensuelle examine les changements majeurs et l’obsolescence.")
    add_table(doc, ["Thème", "Sources", "Fréquence", "Action"], [
        ("Dépendances", "Dependabot, npm audit, GitHub Advisories", "Continue / semaine", "Ticket selon criticité"),
        ("Next.js / React", "Release notes et documentation officielles", "Mensuelle", "Prototype et test de migration"),
        ("Prisma / MySQL", "Changelog Prisma et documentation MySQL", "Mensuelle", "Migration sur base isolée"),
        ("Sécurité web", "OWASP, ANSSI, CERT-FR", "Hebdomadaire", "Comparer à la matrice de risques"),
        ("Données", "CNIL et référentiels HDS", "Mensuelle", "Mettre à jour politiques et contrats"),
        ("Exploitation", "Dokploy, Cloudflare, Node.js", "Mensuelle", "Qualifier versions et fin de support"),
    ], widths=[3.0, 6.1, 3.2, 3.3], font_size=7.8)
    add_heading(doc, "9.2. Registre initial", 2)
    add_table(doc, ["Date", "Signal", "Impact", "Décision / preuve"], [
        ("13/08/2026", "5 vulnérabilités npm en production", "Critique", "Issue #32, montée de versions, audit à 0"),
        ("13/08/2026", "Rate-limit mémoire", "Élevé", "Compteurs MySQL partagés et clés HMAC"),
        ("13/08/2026", "CSP unsafe-eval/inline", "Élevé", "Nonce par requête ; scripts stricts"),
        ("13/08/2026", "Migration d’une base existante", "Élevé", "Sauvegarde, db push unique, baseline"),
        ("13/08/2026", "Données de diagnostic sans HDS", "Critique", "Historisation bloquée en recette"),
    ], widths=[2.6, 6.2, 2.3, 4.5], font_size=7.8)
    add_para(doc, "Chaque entrée future référence une source, un responsable, une date de réévaluation et, si nécessaire, un ticket. Une information sans impact documenté est archivée sans modifier la production.", italic=True, size=8.7)
    page_break(doc)

    # Page 12
    add_heading(doc, "10. Stratégie de sécurisation", 1)
    add_heading(doc, "10.1. Méthode d’analyse", 2)
    add_para(doc, "L’analyse combine inventaire des actifs, menaces, vulnérabilités, probabilité et impact. La criticité initiale est calculée avant contrôle ; la criticité résiduelle est réévaluée après traitement. Échelles : probabilité P de 1 (rare) à 4 (très probable), impact I de 1 (mineur) à 4 (critique), score C = P × I.")
    add_table(doc, ["Score", "Niveau", "Traitement"], [
        ("1–3", "Faible", "Accepter et surveiller"),
        ("4–7", "Modéré", "Planifier et suivre"),
        ("8–11", "Élevé", "Corriger prioritairement"),
        ("12–16", "Critique", "Bloquer ou confiner immédiatement"),
    ], widths=[3.2, 4.2, 8.2], font_size=8.3)
    add_heading(doc, "10.2. Actifs et objectifs", 2)
    add_table(doc, ["Actif", "Confidentialité", "Intégrité", "Disponibilité"], [
        ("Comptes et sessions", "Élevée", "Élevée", "Modérée"),
        ("Résultats de diagnostic", "Critique", "Critique", "Modérée"),
        ("Contenus publics", "Faible", "Élevée", "Élevée"),
        ("Secrets et configuration", "Critique", "Critique", "Élevée"),
        ("Journaux d’audit", "Élevée", "Élevée", "Modérée"),
        ("Code et pipeline", "Modérée", "Élevée", "Élevée"),
    ], widths=[5.6, 3.3, 3.3, 3.4], font_size=8.0)
    add_heading(doc, "10.3. Défense en profondeur", 2)
    add_bullets(doc, [
        "Prévention : validation Zod, contrôle des rôles, Prisma, CSP avec nonce, CSRF, secrets séparés.",
        "Détection : audit applicatif pseudonymisé, npm audit, Dependabot, CI et surveillance Dokploy.",
        "Réduction : rate-limit partagé, stockage HDS bloqué, minimisation et purge par dernière activité.",
        "Récupération : sauvegardes chiffrées, migrations versionnées, rollback et restauration isolée.",
    ], size=9.0)
    page_break(doc)

    # Page 13
    add_heading(doc, "11. Matrice des risques et plan d’actions", 1)
    add_heading(doc, "11.1. Risques techniques", 2)
    risks1 = [
        ("R1", "Dépendance vulnérable", "4", "4", "16 C", "Versions corrigées, audit CI", "1×4=4 M"),
        ("R2", "Force brute / saturation", "4", "3", "12 C", "Rate-limit MySQL, HMAC, audit", "2×3=6 M"),
        ("R3", "XSS via contenu/script", "3", "4", "12 C", "CSP nonce, validation, React", "1×4=4 M"),
        ("R4", "Requête CSRF mutative", "3", "3", "9 E", "Origin/Sec-Fetch-Site, cookies", "1×3=3 F"),
        ("R5", "Élévation de privilège", "3", "4", "12 C", "JWT, proxy, contrôle API serveur", "1×4=4 M"),
        ("R6", "Injection / ID falsifié", "3", "4", "12 C", "Prisma, Zod, appartenance questions", "1×4=4 M"),
    ]
    add_table(doc, ["ID", "Risque", "P", "I", "Initial", "Traitement", "Résiduel"], risks1,
              widths=[1.0, 3.5, 0.7, 0.7, 1.5, 6.1, 2.1], font_size=7.0)
    add_heading(doc, "11.2. Lecture", 2)
    add_para(doc, "Le risque initial est établi dans le contexte d’une application publique. Les traitements réduisent principalement la probabilité ; l’impact intrinsèque d’une fuite de données de santé ou de secrets reste élevé. Un risque résiduel modéré n’est pas considéré comme supprimé : il doit conserver un propriétaire, une surveillance et une date de revue.")
    add_heading(doc, "11.3. Vérifications associées", 2)
    add_table(doc, ["Contrôle", "Preuve"], [
        ("Dépendances", "npm audit à zéro et Dependabot activé"),
        ("CSP", "En-tête avec nonce différent par requête ; scripts correspondants"),
        ("CSRF", "Requête cross-site mutative rejetée en 403"),
        ("Diagnostic", "Questions étrangères/dupliquées rejetées"),
        ("Autorisation", "Tests 401/403 et contrôle des rôles côté serveur"),
        ("Rate-limit", "Compteurs persistants partagés entre instances"),
    ], widths=[4.2, 11.4], font_size=8.1)
    add_para(doc, "Légende : F = faible, M = modéré, E = élevé, C = critique.", italic=True, size=8.2)
    page_break(doc)

    # Page 14
    add_heading(doc, "11. Matrice des risques et plan d’actions (suite)", 1)
    add_heading(doc, "11.4. Données et exploitation", 2)
    risks2 = [
        ("R7", "Données santé sans HDS", "4", "4", "16 C", "Blocage de l’historisation", "1×4=4 M"),
        ("R8", "Fuite via journaux", "3", "4", "12 C", "HMAC IP/compte, métadonnées minimales", "1×4=4 M"),
        ("R9", "Perte ou corruption MySQL", "3", "4", "12 C", "Sauvegarde chiffrée, baseline", "2×4=8 E"),
        ("R10", "Secrets exposés", "3", "4", "12 C", ".gitignore, variables Dokploy, rotation", "1×4=4 M"),
        ("R11", "Indisponibilité VPS/tiers", "3", "3", "9 E", "Monitoring, rollback, services séparés", "2×3=6 M"),
        ("R12", "Purge inadaptée", "3", "3", "9 E", "Dernière activité et procédure", "1×3=3 F"),
        ("R13", "Chaîne CI compromise", "2", "4", "8 E", "Permissions minimales, MAJ actions", "1×4=4 M"),
    ]
    add_table(doc, ["ID", "Risque", "P", "I", "Initial", "Traitement", "Résiduel"], risks2,
              widths=[1.0, 3.6, 0.7, 0.7, 1.5, 6.0, 2.1], font_size=6.9)
    add_heading(doc, "11.5. Plan d’actions priorisé", 2)
    add_table(doc, ["Priorité", "Action", "Responsable", "Échéance / état"], [
        ("P1", "Maintenir HDS_COMPLIANT_STORAGE à 0", "Responsable", "En place"),
        ("P1", "Restaurer une sauvegarde sur base isolée", "Exploitant", "Mensuel — à prouver"),
        ("P2", "Sortir les secrets du contexte de build Nixpacks", "Exploitant", "Avant production réelle"),
        ("P2", "Ajouter supervision et alertes externes", "Exploitant", "Avant production réelle"),
        ("P2", "Créer un tag et une GitHub Release", "Responsable", "Avant soutenance"),
        ("P3", "Résoudre les avertissements lint restants", "Développeur", "Prochain cycle"),
    ], widths=[2.0, 7.2, 3.0, 3.4], font_size=7.8)
    add_para(doc, "Le risque R9 demeure élevé après contrôle tant qu’un test complet de restauration n’a pas été exécuté. Cette transparence évite de confondre une archive déchiffrable avec une reprise réellement éprouvée.", italic=True, size=8.6)
    page_break(doc)

    # Page 15
    add_heading(doc, "12. Données personnelles et RGPD", 1)
    add_heading(doc, "12.1. Catégories et finalités", 2)
    add_table(doc, ["Donnée", "Finalité", "Protection / durée"], [
        ("Identité et email", "Compte et authentification", "Accès restreint ; suppression du compte"),
        ("Mot de passe", "Authentification", "Hash bcrypt avec sel ; jamais exporté"),
        ("Consentement", "Prouver le choix utilisateur", "Horodatage/état associé au compte"),
        ("Score diagnostic", "Résultat et historique autorisé", "Non conservé dans la recette non HDS"),
        ("Journaux", "Sécurité et investigation", "HMAC, métadonnées minimales, purge"),
        ("Préférences UI", "Accessibilité et cookies", "Stockage local ; pas de traceur tiers"),
    ], widths=[3.5, 5.4, 6.7], font_size=7.9)
    add_heading(doc, "12.2. Droits et outils", 2)
    add_bullets(doc, [
        "Information : politique de confidentialité, finalités et avertissement sur le caractère indicatif du diagnostic.",
        "Consentement explicite avant le traitement du questionnaire et lors de l’inscription.",
        "Accès et rectification : consultation et modification du profil authentifié.",
        "Portabilité : export JSON des données du compte, sans mot de passe.",
        "Effacement : suppression du compte et des relations en cascade.",
        "Minimisation : absence de traqueur publicitaire ; journaux pseudonymisés ; score exclu des logs.",
        "Conservation : purge fondée sur la dernière activité et politique opérationnelle documentée.",
    ], size=8.8)
    add_heading(doc, "12.3. Sous-traitants et données sensibles", 2)
    add_para(doc, "Cloudflare, l’hébergeur du VPS et Alwaysdata doivent être recensés avec leur rôle, localisation, garanties et clauses contractuelles. L’environnement de recette ne revendique aucune conformité HDS. La variable de garde n’est pas une certification : l’activation exigerait la vérification du certificat HDS, du périmètre couvert, du contrat, des sauvegardes et de la procédure d’incident.")
    add_heading(doc, "12.4. Écart identifié", 2)
    add_para(doc, "Le message de consentement doit rester cohérent avec le mode de recette : lorsque le stockage est désactivé, l’utilisateur doit être informé que le score est calculé sans être ajouté à son historique. Cette adaptation est enregistrée comme amélioration avant toute recette formelle.", italic=True, size=8.7)
    page_break(doc)

    # Page 16
    add_heading(doc, "13. Gestion de crise", 1)
    add_heading(doc, "13.1. Déclenchement", 2)
    add_para(doc, "Une crise est déclarée lorsqu’un incident menace la confidentialité, l’intégrité ou la disponibilité de manière critique : compromission d’un compte administrateur, exposition de secrets, fuite de données, altération de la base ou indisponibilité durable. Le premier objectif est de limiter l’impact sans détruire les preuves.")
    add_table(doc, ["Phase", "Décision / action", "Responsable"], [
        ("Détection", "Horodater, conserver les logs, ouvrir un canal restreint", "Détecteur"),
        ("Qualification", "Périmètre, données, comptes, criticité, chronologie", "Responsable incident"),
        ("Confinement", "Bloquer route/compte, isoler service, révoquer sessions", "Exploitant"),
        ("Éradication", "Corriger, renouveler secrets, supprimer persistance", "Développeur"),
        ("Rétablissement", "Restaurer, migrer, tester, remettre progressivement", "Exploitant"),
        ("Communication", "Informer les parties selon impact et obligations", "Responsable"),
        ("Clôture", "Surveiller, établir le REX et le plan d’action", "Responsable"),
    ], widths=[2.7, 9.7, 3.2], font_size=7.9)
    add_heading(doc, "13.2. Escalade", 2)
    add_table(doc, ["Niveau", "Exemple", "Délai", "Canal"], [
        ("P1", "Fuite, compromission, service critique indisponible", "Immédiat", "Téléphone + canal privé"),
        ("P2", "Faille exploitable ou fonction majeure", "Journée", "Ticket restreint + message"),
        ("P3/P4", "Incident limité ou amélioration", "Cycle normal", "GitHub Issue"),
    ], widths=[2.1, 7.3, 2.8, 3.4], font_size=8.0)
    add_heading(doc, "13.3. Communication", 2)
    add_para(doc, "Le message initial contient les faits confirmés, l’heure, le périmètre présumé, les mesures prises, l’impact utilisateur et la prochaine échéance d’information. Aucun diagnostic, email, IP brute, secret ou hypothèse non vérifiée n’est publié. Les notifications réglementaires sont préparées avec le responsable compétent selon la nature et le risque de la violation.")
    add_para(doc, "Après reprise : chronologie, cause racine, efficacité des contrôles, décisions, personnes informées, données concernées, temps de rétablissement et actions datées.", italic=True, size=8.7)
    page_break(doc)

    # Page 17
    add_heading(doc, "14. Bonnes pratiques et validation", 1)
    add_heading(doc, "14.1. Développement sécurisé", 2)
    add_table(doc, ["Pratique", "Application dans CESIZen"], [
        ("Validation", "Schémas Zod et contrôles métier côté serveur"),
        ("Accès", "Session, rôles et propriétaire vérifiés sur les routes"),
        ("Données", "Prisma paramétré, sélections minimales, cascade maîtrisée"),
        ("Secrets", ".env ignoré, exemple sans valeur, clés distinctes"),
        ("Navigateur", "CSP nonce, anti-framing, HTTPS/HSTS, politique de permissions"),
        ("Traçabilité", "Audit pseudonymisé et métadonnées filtrées"),
        ("Dépendances", "Lockfile, npm audit, Dependabot et mises à jour testées"),
        ("Documentation", "README, tests, sécurité, sauvegarde et maintenance"),
    ], widths=[4.1, 11.5], font_size=8.1)
    add_heading(doc, "14.2. Pyramide de tests", 2)
    add_table(doc, ["Niveau", "Périmètre", "Résultat"], [
        ("Unitaire", "Validation, calcul, sécurité, services", "Automatisé"),
        ("Intégration", "Routes API, auth, RGPD, diagnostics", "Automatisé avec mocks maîtrisés"),
        ("E2E", "Parcours navigateur et accessibilité", "Playwright, exécution ciblée"),
        ("Build", "Compilation et rendu Next.js", "Réussi"),
        ("Smoke", "Accueil, auth, contenu, diagnostic, admin", "Après déploiement"),
    ], widths=[3.0, 8.0, 4.6], font_size=8.2)
    add_heading(doc, "14.3. Résultats", 2)
    add_bullets(doc, [
        "30 fichiers de tests Vitest, 180 tests réussis.",
        "TypeScript sans erreur ; ESLint sans erreur bloquante, 15 avertissements suivis.",
        "Build de production réussi avec Next.js 16.3.0.",
        "Audit npm : zéro vulnérabilité connue au 13 août 2026.",
        "CI GitHub Actions verte sur le commit déployé.",
        "Contrôle navigateur : aucune erreur d’hydratation ou de nonce après correctif.",
    ], size=8.9)
    page_break(doc)

    # Page 18
    add_heading(doc, "15. Démonstration et éléments de preuve", 1)
    add_heading(doc, "15.1. Scénario de démonstration (6 minutes)", 2)
    add_numbered(doc, [
        "Ouvrir GitHub : montrer main/develop, leurs protections, une pull request et les fichiers exclus.",
        "Ouvrir l’issue #32 : criticité, critères, correction, tests et retour arrière.",
        "Ouvrir GitHub Actions : audit, migration MySQL, lint, types, tests et build verts.",
        "Ouvrir Dokploy : commit déployé, variables masquées, logs et état du conteneur.",
        "Ouvrir cesizen.optihent.fr : HTTPS, informations, diagnostic anonyme et contrôle admin.",
        "Expliquer HDS_COMPLIANT_STORAGE=0 : calcul possible, historique volontairement bloqué.",
        "Conclure par la sauvegarde, la restauration, la veille et la réponse à incident.",
    ], size=8.8)
    add_heading(doc, "15.2. Preuves consultables", 2)
    add_table(doc, ["Preuve", "Emplacement"], [
        ("Code et historique", "github.com/MarinCD/CesiZen"),
        ("Recette HTTPS", "https://cesizen.optihent.fr"),
        ("CI", "GitHub → Actions → CI → commit 9ce6e4e"),
        ("Ticket sécurité", "GitHub Issue #32"),
        ("Déploiement", "GitHub Issue #31 et journaux Dokploy"),
        ("Migrations", "prisma/migrations/20260813170000_baseline"),
        ("Tests", "docs/TESTING.md et docs/RUNNING_TESTS.md"),
        ("Maintenance", "docs/MAINTENANCE.md"),
        ("Sauvegarde", "docs/BACKUP_RESTORE.md"),
        ("Sécurité", "docs/SECURITY_HARDENING.md et SECURITY.md"),
    ], widths=[4.0, 11.6], font_size=7.7)
    add_heading(doc, "15.3. Critères de réussite de la soutenance", 2)
    add_para(doc, "La démonstration est préparée avec des onglets ouverts et un compte de démonstration sans donnée réelle. Les secrets sont masqués. Un export PDF du dossier et des captures hors ligne sont conservés pour poursuivre la présentation en cas d’indisponibilité réseau.")
    page_break(doc)

    # Page 19
    add_heading(doc, "16. Limites, évolutions et conclusion", 1)
    add_heading(doc, "16.1. Limites connues", 2)
    add_table(doc, ["Limite", "Conséquence", "Traitement prévu"], [
        ("Hébergement non qualifié HDS", "Pas d’historique santé en recette", "Choisir et contractualiser un hébergeur qualifié"),
        ("Restauration complète non éprouvée", "RTO non confirmé", "Test mensuel sur base isolée"),
        ("Monitoring externe minimal", "Détection tardive possible", "Sonde, alertes et centralisation"),
        ("Build Nixpacks avec variables", "Avertissements sur les secrets", "Image multi-stage et secrets runtime"),
        ("15 avertissements lint", "Dette de qualité mineure", "Ticket et correction incrémentale"),
        ("Pas de mot de passe oublié", "Réinitialisation hors profil absente", "Jeton court, email et invalidation"),
        ("Résultat diagnostic codé", "Configuration admin partielle", "Modèle de seuils versionné"),
    ], widths=[4.3, 5.4, 5.9], font_size=7.7)
    add_heading(doc, "16.2. Feuille de route", 2)
    add_bullets(doc, [
        "Avant soutenance : clôturer les issues #31/#32 avec preuves, créer un tag/release et préparer les captures.",
        "Court terme : test de restauration, alertes externes, suppression des avertissements lint et message de consentement conditionnel.",
        "Moyen terme : mot de passe oublié, résultats administrables et automatisation de promotion contrôlée.",
        "Avant production réelle : qualification HDS, contrats, analyse d’impact, supervision, astreinte et exercice de crise.",
    ], size=8.8)
    add_heading(doc, "16.3. Conclusion", 2)
    add_para(doc, "CESIZen dispose d’une chaîne démontrable allant du code versionné à une recette HTTPS : GitHub structure les sources et demandes, GitHub Actions automatise la qualité, Dokploy orchestre l’application et Prisma maîtrise les évolutions de schéma. Les corrections de sécurité traitent les risques prioritaires et les données de diagnostic ne sont pas persistées sans garantie HDS.")
    add_para(doc, "La valeur du dispositif repose enfin sur sa gouvernance : tickets qualifiés, preuves, veille, sauvegardes, retour arrière et gestion de crise. Les limites restantes sont explicites, priorisées et associées à un responsable ou à une condition d’ouverture de la production.")
    page_break(doc)

    # Page 20
    add_heading(doc, "17. Glossaire et références", 1)
    add_heading(doc, "17.1. Glossaire", 2)
    add_table(doc, ["Terme", "Définition"], [
        ("CI/CD", "Intégration continue et livraison/déploiement continu."),
        ("CSP", "Politique navigateur limitant les sources autorisées."),
        ("CSRF", "Requête forgée utilisant la session d’une victime."),
        ("HDS", "Hébergement de données de santé soumis à certification."),
        ("HMAC", "Code d’authentification produisant ici un pseudonyme stable."),
        ("RPO", "Perte de données maximale admissible exprimée dans le temps."),
        ("RTO", "Temps cible de rétablissement d’un service."),
        ("SLA", "Engagement de délai ou de niveau de service."),
        ("Smoke test", "Contrôle court des fonctions essentielles après publication."),
    ], widths=[3.0, 12.6], font_size=8.0)
    add_heading(doc, "17.2. Références", 2)
    refs = [
        "Consignes INFCDAAL3 — Déployer et sécuriser les applications informatiques, CESI, V1 du 07/06/2024.",
        "Grille d’évaluation Bloc 3 — Déployer et sécuriser les applications informatiques.",
        "OWASP Application Security Verification Standard et OWASP Cheat Sheet Series — owasp.org.",
        "CNIL — sécurité des données personnelles, violations de données et droits des personnes — cnil.fr.",
        "Agence du Numérique en Santé — certification des hébergeurs de données de santé — esante.gouv.fr.",
        "Next.js — documentation du déploiement et Content Security Policy — nextjs.org/docs.",
        "Prisma — Prisma Migrate en production — prisma.io/docs.",
        "Dokploy — Applications, domaines, variables et mise en production — docs.dokploy.com.",
        "GitHub — Actions, Issues, Dependabot et Private Vulnerability Reporting — docs.github.com.",
    ]
    add_bullets(doc, refs, size=8.2)
    add_heading(doc, "17.3. Documents du dépôt", 2)
    add_para(doc, "README.md ; docs/MAINTENANCE.md ; docs/BACKUP_RESTORE.md ; docs/SECURITY_HARDENING.md ; docs/TESTING.md ; SECURITY.md ; .github/workflows/ci.yml ; prisma/schema.prisma.", size=8.6)
    add_para(doc, "Fin du dossier — version 1.1 du 13 août 2026", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, before=14, size=10)

    doc.core_properties.title = "CESIZen — Bloc 3 : déploiement, maintenance et sécurisation"
    doc.core_properties.subject = "Dossier d’évaluation INFCDAAL3"
    doc.core_properties.author = "Marin Cadro"
    doc.core_properties.keywords = "CESIZen, déploiement, maintenance, sécurité, CESI, CDA"
    doc.core_properties.comments = "Généré à partir du style Documentation_Technique_CESIZen_v3."
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
